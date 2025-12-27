from docx import Document
import os

# Create data directory
if not os.path.exists("data"):
    os.makedirs("data")

# 1. Create JD
jd_text = """
Job Title: Senior Python Developer
Location: Remote

We are looking for a Senior Python Developer to join our AI team.

Responsibilities:
- Build scalable backend APIs using FastAPI.
- specific experience with LangChain and vector databases.
- Deploy models using Docker and Kubernetes.

Requirements:
- 5+ years of experience in Python.
- Strong knowledge of FastAPI, Pydantic, and SQLAlchemy.
- Experience with LLMs (Ollama, OpenAI API).
- Cloud experience (AWS or GCP).
- Bachelor's degree in Computer Science or related field.

Nice to have:
- Open source contributions.
- Experience with React/Next.js.
"""

with open("data/jd.txt", "w") as f:
    f.write(jd_text.strip())

# 2. Create Resume DOCX
doc = Document()
doc.add_heading('John Doe', 0)
doc.add_paragraph('john.doe@email.com | (555) 123-4567')

doc.add_heading('Summary', level=1)
doc.add_paragraph('Experienced Software Engineer with 6 years of expertise in building scalable web applications and AI systems. Passionate about LLMs.')

doc.add_heading('Experience', level=1)
doc.add_paragraph('Senior Software Engineer - Tech Corp (2020 - Present)', style='List Bullet')
doc.add_paragraph('Developed microservices using Python and FastAPI.')
doc.add_paragraph('Integrated LLM features using LangChain and ChromaDB.')
doc.add_paragraph('Managed deployments on AWS EKS.')

doc.add_paragraph('Software Developer - StartUp Inc (2017 - 2020)', style='List Bullet')
doc.add_paragraph('Built REST APIs using Django.')
doc.add_paragraph('Worked with PostgreSQL and Redis.')

doc.add_heading('Education', level=1)
doc.add_paragraph('Bachelor of Science in Computer Science - University of Tech (2013-2017)')

doc.add_heading('Skills', level=1)
doc.add_paragraph('Python, FastAPI, Django, Docker, Kubernetes, AWS, LangChain, SQL, Git')

doc.save("data/resume_sample.docx")

print("Sample data created in data/ folder.")
